import os
import zipfile
import re
from typing import Optional
from xml.etree import ElementTree as ET


def _extract_odt_text(file_path: str) -> str:
    with zipfile.ZipFile(file_path) as archive:
        content_xml = archive.read("content.xml")

    root = ET.fromstring(content_xml)
    chunks: list[str] = []

    def traverse(node):
        tag_name = node.tag.split('}')[-1] if '}' in node.tag else node.tag
        is_block = tag_name in ['p', 'h', 'list-item']

        if node.text:
            chunks.append(node.text)

        for child in node:
            traverse(child)

        if tag_name == 'table-cell':
            chunks.append(" | ")
        elif tag_name == 'table-row':
            chunks.append("\n")

        if is_block:
            chunks.append("\n")

        if node.tail:
            chunks.append(node.tail)

    traverse(root)

    text = "".join(chunks)
    text = re.sub(r' +', ' ', text)
    text = text.replace(' | \n', '\n')
    text = re.sub(r'\n+', '\n', text)
    return text.strip()


def extract_text_from_file(file_path: str, file_type: str) -> Optional[str]:
    """
    根據副檔名解析文檔並萃取文字
    """
    if not os.path.exists(file_path):
        return None

    try:
        if file_type == "pdf":
            import PyPDF2

            text = ""
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n"
            return text.strip()

        elif file_type == "docx":
            from docx import Document

            doc = Document(file_path)
            text_blocks = []
            
            # 1. 抓取所有段落文字
            for para in doc.paragraphs:
                cleaned = para.text.strip()
                if cleaned:
                    text_blocks.append(cleaned)
                    
            # 2. 抓取所有表格文字 (許多履歷與排版文件會使用表格)
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cleaned_cell = cell.text.strip().replace('\n', ' ')
                        if cleaned_cell:
                            row_data.append(cleaned_cell)
                    if row_data:
                        text_blocks.append(" | ".join(row_data))
                        
            text = "\n".join(text_blocks)
            return text.strip()

        elif file_type == "doc":
            import subprocess
            # 使用系統安裝的 antiword 解析 legacy .doc (二進制格式)
            try:
                result = subprocess.run(['antiword', file_path], capture_output=True, text=True, check=True)
                return result.stdout.strip()
            except FileNotFoundError:
                raise RuntimeError("System dependency 'antiword' is not installed. Failed to parse legacy .doc file.")
            except subprocess.CalledProcessError as e:
                raise RuntimeError(f"antiword execution failed. Output: {e.stderr}")

        elif file_type in ["xlsx", "xls", "csv"]:
            import pandas as pd

            if file_type == "csv":
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            # 將表格轉為字串表示
            return df.to_string(index=False)

        elif file_type == "txt":
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read().strip()

        elif file_type == "odt":
            return _extract_odt_text(file_path)

        else:
            # 不支援的格式
            raise ValueError(f"Unsupported file type: {file_type}")

    except Exception as e:
        print(f"Error parsing file {file_path}: {e}")
        return None
